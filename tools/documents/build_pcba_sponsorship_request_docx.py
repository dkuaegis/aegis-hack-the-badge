from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SPONSORSHIP_DIR = ROOT / "docs" / "rev3" / "sponsorship" / "jlcpcb"
INPUT = SPONSORSHIP_DIR / "hack-the-badge-jlcpcb-support-brief.docx"
OUT = SPONSORSHIP_DIR / "Hack_The_Badge_JLCPCB_PCBA_Sponsorship_Request.docx"

LOGO = ROOT / "assets" / "brand" / "aegis" / "black-white-ring.png"
RELEASE_DIR = ROOT / "hardware" / "releases" / "rev3" / "jlcpcb"
PCB_ISO = RELEASE_DIR / "preview" / "final_preview_iso.png"
PCB_TOP = RELEASE_DIR / "preview" / "final_preview_top.png"

GERBER = RELEASE_DIR / "upload" / "hacking_badge_v3_jlcpcb_gerbers.zip"
BOM = RELEASE_DIR / "upload" / "hacking_badge_v3_jlcpcb_bom.csv"
CPL = RELEASE_DIR / "upload" / "hacking_badge_v3_jlcpcb_cpl.csv"
MANUAL = RELEASE_DIR / "assembly" / "hacking_badge_v3_manual_assembly.csv"
VALIDATION = RELEASE_DIR / "reports" / "manufacturing_validation.txt"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
MID_FILL = "E8EEF5"
BORDER = "B7C4D4"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="2E74B5", size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167


def add_para(doc, text="", style=None, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_font(run, size=11, color=INK)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_font(run, size=11, color=INK)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_FILL, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_fill(cell, header_fill)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_font(r, size=9.5, color=INK, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for r in p.runs:
                    set_font(r, size=font_size, color=INK)

    set_table_width(table, widths)
    set_table_borders(table)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED, italic=True)


def add_source_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    lr = p.add_run(f"{label}: ")
    set_font(lr, size=10.5, color=INK, bold=True)
    vr = p.add_run(value)
    set_font(vr, size=10.5, color=BLUE if value.startswith("http") else INK)


def file_status(path):
    return "Available for immediate review" if path.exists() else "To be provided by [DATE]"


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PCBA SPONSORSHIP REQUEST")
    set_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Hack The Badge")
    set_font(r, size=25, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Request for JLCPCB sponsorship of approximately 10 fully assembled PCBA units for MSG CTF")
    set_font(r, size=12, color=MUTED)

    meta = [
        ("Requester", "Aegis Security Team / MSG CTF"),
        ("Request type", "PCB fabrication, component procurement, SMT assembly, and delivery support"),
        ("Requested quantity", "Approximately 10 completed units"),
        ("Delivery destination", "Republic of Korea"),
        ("Schedule request", "Preferably within 7 days; absolute maximum 14 days after production approval"),
        ("Design status", "Circuit and PCB design completed; production files are available for review"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_font(vr, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule)


def add_image_pair(doc):
    add_heading(doc, "Project Visuals", 2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_width(table, [3200, 5860])
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    c1, c2 = table.rows[0].cells
    if LOGO.exists():
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.55))
    if PCB_ISO.exists():
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(PCB_ISO), width=Inches(3.3))
    add_caption(doc, "Visual references: Aegis badge artwork and current Hack The Badge PCB render.")


def main():
    # The existing DOCX is treated as the source artifact for the visual direction;
    # this file intentionally writes a new output DOCX with updated sponsorship content.
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "Hack The Badge - JLCPCB PCBA Sponsorship Request"
    for run in header.runs:
        set_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "PCBA sponsorship request"
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)

    add_masthead(doc)

    add_heading(doc, "Request Summary", 1)
    add_para(
        doc,
        "Hack The Badge is an interactive cybersecurity badge project operated by the Aegis Security Team for the MSG CTF finals booth. "
        "Participants connect to the badge over USB Serial and solve cybersecurity challenges embedded in the badge experience.",
    )
    add_para(
        doc,
        "We are requesting JLCPCB sponsorship support for approximately 10 fully assembled PCBA units. In this document, "
        "completed units means assembled PCBA boards ready for project-side firmware loading and event preparation. "
        "This request does not ask JLCPCB to provide an enclosure, retail packaging, or firmware programming unless those services are separately offered and confirmed.",
    )

    add_table(
        doc,
        ["Item", "Request / current status"],
        [
            ("Requested support", "PCB fabrication, parts, component procurement, SMT assembly, and delivery support as available from JLCPCB."),
            ("Quantity", "Approximately 10 completed PCBA units."),
            ("Timing", "Please confirm whether completion and delivery to South Korea is possible within 7 days if feasible, or within 14 days at the latest."),
            ("Design status", "Circuit and PCB design are completed. This is a production feasibility and urgent schedule request, not a pre-design review."),
            ("Destination", "Republic of Korea."),
        ],
        [2100, 7260],
        MID_FILL,
    )

    add_image_pair(doc)

    add_heading(doc, "Production Files Available For Review", 1)
    add_para(
        doc,
        "The current project folder includes JLCPCB-oriented manufacturing outputs. These files can be provided for immediate review. "
        "The quoted quantity should be updated for the requested sponsorship build of approximately 10 units.",
    )
    add_table(
        doc,
        ["File / artifact", "Status", "Notes"],
        [
            ("Gerber ZIP", file_status(GERBER), "JLCPCB fabrication package is present in the project folder."),
            ("JLCPCB BOM CSV", file_status(BOM), "Current SMT assembly BOM is present; live stock and pricing should be rechecked by JLCPCB."),
            ("JLCPCB CPL CSV", file_status(CPL), "Current component placement file is present for review."),
            ("Manufacturing validation", file_status(VALIDATION), "Local validation report is present and records the available package state."),
            ("Manual assembly list", file_status(MANUAL), "OLED1 is listed separately in the current package; please confirm whether JLCPCB can support it for completed PCBA units."),
            ("Final sponsor submission date", "To be confirmed", "Use [DATE] as the placeholder until the actual submission date is set."),
        ],
        [2350, 2600, 4410],
        LIGHT_FILL,
        font_size=8.9,
    )

    add_heading(doc, "Hardware Scope", 1)
    add_para(
        doc,
        "The badge is based on an ESP32-S3-WROOM class MCU and is designed as a USB-powered interactive CTF device. "
        "The main participant interface is USB Serial, with additional intentional challenge pads for UART/GPIO interaction.",
    )
    add_table(
        doc,
        ["Subsystem", "Included in the design"],
        [
            ("Controller", "ESP32-S3-WROOM class module."),
            ("USB and power", "USB-C device connector and required power/protection circuitry."),
            ("User interface", "OLED display, tactile switches, and LEDs."),
            ("Challenge interface", "UART/GPIO challenge pads and required protection components."),
            ("Assembly target", "Fully assembled PCBA boards. Please confirm whether OLED sourcing/placement can be included by JLCPCB."),
        ],
        [2400, 6960],
        MID_FILL,
    )

    add_heading(doc, "Requested Sponsorship Scope Confirmation", 1)
    add_para(
        doc,
        "Please confirm which parts of the production and delivery scope JLCPCB may be able to sponsor or discount. "
        "Partial sponsorship is still useful if full coverage is not available.",
    )
    add_table(
        doc,
        ["Scope area", "Confirmation requested from JLCPCB"],
        [
            ("PCB fabrication", "Can JLCPCB sponsor or discount the PCB fabrication cost for approximately 10 units?"),
            ("Component cost", "Can JLCPCB sponsor or discount the electronic component cost?"),
            ("Component procurement", "Can JLCPCB procure all listed PCBA components, or identify parts requiring substitution/approval?"),
            ("SMT assembly", "Can JLCPCB sponsor or discount SMT assembly for the requested quantity?"),
            ("OLED / non-SMT handling", "Can JLCPCB support the OLED as part of completed PCBA units, or should it remain outside JLCPCB assembly?"),
            ("Shipping", "Can JLCPCB sponsor, discount, or expedite shipping of the completed PCBA units to South Korea?"),
        ],
        [2600, 6760],
        LIGHT_FILL,
    )

    doc.add_page_break()
    add_heading(doc, "Urgent Schedule Confirmation", 1)
    add_para(
        doc,
        "The project is intended for use at the MSG CTF finals booth. Please confirm the fastest realistic schedule after file review and production approval. "
        "We are specifically asking whether completed PCBA units can be fabricated, assembled, and shipped to South Korea within the following windows.",
    )
    add_table(
        doc,
        ["Schedule window", "Requested confirmation"],
        [
            ("Preferred", "Completion and delivery to South Korea within 7 days if feasible."),
            ("Maximum acceptable", "Completion and delivery to South Korea within 14 days at the latest."),
            ("If not feasible", "Please provide the earliest realistic production and delivery schedule and identify the bottleneck."),
            ("Decision date", "Please use [DATE] as a placeholder until the actual sponsorship decision date is confirmed."),
        ],
        [2500, 6860],
        MID_FILL,
    )

    add_heading(doc, "Review Questions For JLCPCB", 1)
    for item in [
        "Can JLCPCB support approximately 10 completed PCBA units for this Hack The Badge project?",
        "Can the available Gerber, BOM, and CPL files be reviewed immediately for production feasibility?",
        "Which cost areas can JLCPCB sponsor or discount: PCB fabrication, parts, procurement, SMT assembly, or shipping?",
        "Is the 7-day preferred schedule possible, including shipment to South Korea? If not, is 14 days possible?",
        "Are there any components that require substitution, stock confirmation, or special handling before production approval?",
        "Can JLCPCB include the OLED in the completed PCBA scope, or should the OLED remain a separate project-side assembly step?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "References", 1)
    add_source_line(doc, "Existing project background", "https://github.com/Z3r0c0k3/hacking-box")
    add_source_line(doc, "Local manufacturing package", "hardware/releases/rev3/jlcpcb/")
    add_source_line(doc, "Primary files for review", "hacking_badge_v3_jlcpcb_gerbers.zip, hacking_badge_v3_jlcpcb_bom.csv, hacking_badge_v3_jlcpcb_cpl.csv")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
