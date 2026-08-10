from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_jlcpcb_support_docx import (
    BLUE,
    BORDER,
    DARK_BLUE,
    INK,
    LIGHT_FILL,
    LOGO,
    MID_FILL,
    MUTED,
    PCB_PREVIEW,
    ROOT,
    add_bullet,
    add_caption,
    add_heading,
    add_para,
    add_source_link,
    add_table,
    configure_styles,
    paragraph_border_bottom,
    set_font,
)


OUT = ROOT / "docs" / "rev3" / "sponsorship" / "jlcpcb" / "hack-the-badge-jlcpcb-support-brief-ko.docx"


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("JLCPCB SUPPORT BRIEF - KOREAN REVIEW COPY")
    set_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Hack The Badge / Hacking Box V2")
    set_font(r, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("기존 Arduino Hacking Badge 프로젝트를 기반으로 한 Phase 1 하드웨어 요구사항")
    set_font(r, size=12, color=MUTED)

    meta = [
        ("용도", "JLCPCB 제작 및 PCBA 지원 검토용 문서의 한국어 검토본"),
        ("작성 주체", "Aegis / MSG CTF project team"),
        ("작성일", "2026년 8월 4일"),
        ("상태", "요구사항 초안; Phase 1에서는 회로도와 PCB 레이아웃을 만들지 않음"),
        ("기반 프로젝트", "https://github.com/Z3r0c0k3/hacking-box"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_font(vr, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule)


def add_image_block(doc):
    add_heading(doc, "예시 이미지", 2)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.9))
        add_caption(
            doc,
            "그림 1. 배지 정체성 참고용으로 black-white-ring.svg에서 렌더링한 Aegis 링 아트워크.",
        )
    if PCB_PREVIEW.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(PCB_PREVIEW), width=Inches(3.7))
        add_caption(
            doc,
            "그림 2. 기존 로컬 배지 렌더 이미지. 시각 방향 참고용이며, 이 문서가 새 PCB 설계를 릴리스하는 것은 아님.",
        )


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "Hack The Badge / Hacking Box V2 - Korean Review Copy"
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Phase 1 requirements draft"
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)

    add_masthead(doc)

    add_heading(doc, "요약", 1)
    add_para(
        doc,
        "Hack The Badge는 Aegis / MSG CTF 행사 부스에서 사용할 CTF 하드웨어 해킹 배지다. "
        "이번 JLCPCB 지원 요청의 목적은 회로도 작성 전에 하드웨어 콘셉트를 검토하고, "
        "PCBA 가능성과 제조/부품 선정 리스크를 미리 확인하는 것이다.",
    )
    add_para(
        doc,
        "이 V2 작업은 기존 Z3r0c0k3/hacking-box Arduino Hacking Badge 프로젝트를 기반으로 한다. "
        "기존 프로젝트는 Serial shell 형태의 챌린지, 여러 독립 퍼즐, LED 피드백, solve 상태 저장을 제공한다. "
        "V2는 참가자가 보는 문제 풀이 흐름을 최대한 유지하면서, 소량 행사 제작에 적합한 ESP32-S3 기반 PCB로 전환하는 것을 목표로 한다.",
    )

    add_image_block(doc)

    add_heading(doc, "기존 Hacking Box 프로젝트 기반", 1)
    add_para(
        doc,
        "공개 기준 프로젝트는 GitHub의 Z3r0c0k3/hacking-box다. README는 이 프로젝트를 Arduino Hacking Badge로 설명하며, "
        "대형 CTF 결승에서 배포되는 badge hacking 경험을 모방하는 구성을 가진다. 저장소에는 example, old, template badge "
        "펌웨어 폴더가 있고, Arduino IDE, C/C++, EEPROM, PROGMEM을 기술 기반으로 사용한다.",
    )
    for item in [
        "기존 프로젝트는 4개의 독립 퍼즐 슬롯을 설명한다.",
        "각 퍼즐 solve 상태는 LED 피드백으로 표시되며, 모든 퍼즐을 풀었을 때 추가 LED 애니메이션을 수행한다.",
        "Serial Monitor 기반 shell 인터페이스로 게임 시작, 문제 선택, 상태 확인, 진행 상황 초기화, 화면 지우기, 문제 종료를 수행한다.",
        "README는 4800 baud Serial Monitor 설정과 Arduino Uno/Nano 스타일 하드웨어 기준을 문서화한다.",
        "긴 문제 텍스트와 ASCII art는 PROGMEM에 저장하고, solve 상태는 EEPROM에 저장한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "JLCPCB에 요청하려는 지원", 1)
    for item in [
        "회로도와 PCB 레이아웃 전에 제안 하드웨어 구조를 검토한다.",
        "소량 SMT 조립에 적합한 JLCPCB 수급 가능 부품을 추천받는다.",
        "OLED 모듈을 JLCPCB PCBA에 포함할지, 후납땜으로 처리할지 조언받는다.",
        "USB-C, ESD, 입력 보호, 노출 UART/GPIO 패드, 버튼, LED, 선택 부저가 PCBA에 적합한지 검토받는다.",
        "배지형 PCB 외형과 장식용 silkscreen artwork에 대한 DFM 이슈를 확인한다.",
        "Gerber, BOM, CPL 생성 전에 반드시 확정해야 하는 결정 사항을 확인한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "목표 하드웨어 기준안", 1)
    add_table(
        doc,
        ["기능", "Phase 1 기준안", "이유"],
        [
            ("MCU", "ESP32-S3-WROOM 모듈", "Native USB, 선택적 Wi-Fi/BLE, 더 큰 메모리, bare chip 대비 낮은 RF 리스크."),
            ("로직 전압", "3.3 V", "ESP32-S3와 최신 주변 모듈 기준에 맞음."),
            ("USB", "USB-C device port", "전원, Serial console, 펌웨어 업로드, 운영진 복구에 사용."),
            ("디스플레이", "0.96 inch 128x64 I2C OLED", "로고, 상태, 문제 텍스트를 작게 표시 가능."),
            ("입력", "택트 스위치 3개", "Left/OK/Right 조작, 미니게임 입력, 관리자 unlock 조합에 활용 가능."),
            ("출력", "상태 LED 5개", "문제 진행 상태와 solve 여부를 명확히 표시."),
            ("오디오", "승인 시 드라이버가 있는 패시브 부저", "MCU GPIO에 직접 부하를 걸지 않고 효과음 제공."),
            ("챌린지 I/O", "UART와 보호된 GPIO 패드", "참가자를 위한 의도된 하드웨어 해킹 표면."),
        ],
        [1800, 2700, 4860],
        MID_FILL,
    )

    add_heading(doc, "호환성 요구사항", 1)
    add_para(
        doc,
        "V2는 Arduino Nano / ATmega328P 기준 구현과 바이너리 호환되는 대체품이 아니라, ESP32-S3 기준 펌웨어 이식 또는 재작성으로 보아야 한다. "
        "다만 참가자가 체감하는 동작과 문제 풀이 흐름은 최대한 익숙하게 유지하는 것이 목표다.",
    )
    for item in [
        "USB Serial shell 형태의 챌린지 흐름을 유지한다.",
        "가능한 범위에서 문제 선택 모델과 status 명령 흐름을 유지한다.",
        "solve 상태 피드백을 여러 LED에 매핑하고, all-solved 애니메이션을 넣을 여지를 둔다.",
        "EEPROM 기반 상태 저장은 ESP32-S3의 비휘발성 저장소로 대체하되, 사용자에게 보이는 동작은 유지한다.",
        "3.3 V 보드에서 신호를 노출하기 전에 V1이 5 V Arduino 전기적 동작에 의존하는지 검토한다.",
        "USB flashing, EN/reset, BOOT 접근을 통해 운영진 현장 복구가 가능해야 한다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "제안 ESP32-S3 핀맵", 1)
    add_table(
        doc,
        ["기능", "신호", "동작", "메모"],
        [
            ("USB D-", "GPIO19", "Native USB", "ESP32-S3 USB 사용 시 고정에 가까움."),
            ("USB D+", "GPIO20", "Native USB", "ESP32-S3 USB 사용 시 고정에 가까움."),
            ("OLED SDA", "GPIO4", "3.3 V I2C", "최종 OLED pull-up 여부 확인."),
            ("OLED SCL", "GPIO5", "3.3 V I2C", "확인 전까지 display address 0x3C 가정."),
            ("Challenge 0", "GPIO6", "직렬 저항을 거친 3.3 V", "챌린지 설계 승인 후 노출."),
            ("Challenge 1", "GPIO7", "직렬 저항을 거친 3.3 V", "챌린지 설계 승인 후 노출."),
            ("Challenge 2", "GPIO8", "직렬 저항을 거친 3.3 V", "챌린지 설계 승인 후 노출."),
            ("Left / OK / Right", "GPIO9 / GPIO10 / GPIO12", "Active-low inputs", "Pull-up과 debounce 적용."),
            ("Status LEDs", "GPIO13-GPIO17", "Active-high outputs", "부팅 시 low로 초기화."),
            ("Buzzer PWM", "GPIO18", "MOSFET-driven PWM", "부저 승인 시에만 사용."),
            ("Player UART TX/RX", "GPIO43 / GPIO44", "직렬 저항을 거친 3.3 V", "5 V tolerant 아님."),
            ("Boot / reset", "GPIO0 / EN", "운영진 복구", "챌린지 GPIO로 사용하지 않음."),
        ],
        [1800, 1800, 2760, 3000],
        LIGHT_FILL,
    )

    add_heading(doc, "회로도 작성 전 승인 필요 항목", 1)
    add_table(
        doc,
        ["ID", "필요한 결정", "중요한 이유"],
        [
            ("O-001", "정확한 ESP32-S3-WROOM 모듈 variant 확정.", "BOM, 펌웨어 파티션, 수급성, 가격에 영향을 줌."),
            ("O-002", "첫 제작 수량 확정.", "단가, 예비품 전략, 조립 옵션에 영향을 줌."),
            ("O-003", "OLED를 PCBA에 포함할지 후납땜할지 결정.", "BOM/CPL, footprint, 교체 workflow에 영향을 줌."),
            ("O-004", "부저를 포함할지 제거할지 승인.", "피드백은 좋지만 보드 면적과 펌웨어 작업이 증가함."),
            ("O-005", "챌린지 패드 동작과 라벨 확정.", "핵심 하드웨어 해킹 경험을 정의함."),
            ("O-006", "Wi-Fi/BLE 관리 정책 정의.", "관리 기능이 의도치 않은 우회 경로가 되면 안 됨."),
            ("O-007", "USB와 노출 GPIO/UART 보호 수준 결정.", "참가자가 직접 만지는 보드라 전기적 오사용 리스크가 있음."),
            ("O-008", "배지 외형과 silkscreen artwork 승인.", "DFM, 브랜딩, 지원 검토에 영향을 줌."),
        ],
        [1000, 3900, 4460],
        MID_FILL,
    )

    add_heading(doc, "JLCPCB 검토 질문", 1)
    for item in [
        "현재 PCBA 수급성과 비용 기준으로 어떤 ESP32-S3-WROOM 모듈 variant를 추천하는가?",
        "OLED 모듈을 PCBA에 포함하는 것이 좋은가, 아니면 수동 조립으로 처리하는 것이 좋은가?",
        "제안한 USB-C receptacle, tactile switch, LED, passive buzzer가 단면 SMT 조립에 적합한가?",
        "참가자가 직접 만지는 USB 전원 보드에 input fuse/PTC를 넣는 것이 권장되는가?",
        "노출 UART와 challenge pad에 추가 ESD 또는 보호 부품이 필요한가?",
        "방패형 외형과 black/white 장식 silkscreen artwork에 DFM 우려가 있는가?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Phase 1 범위 경계", 1)
    add_para(
        doc,
        "이 문서는 의도적으로 회로도 작성 전 단계에서 멈춘다. 다음 단계는 위의 미결정 사항이 승인되고, "
        "V1 펌웨어 동작이 제안 ESP32-S3 핀맵과 대조된 뒤에 시작하는 것이 적절하다.",
    )
    for item in [
        "이 문서는 새 회로도를 릴리스하지 않는다.",
        "이 문서는 새 PCB 레이아웃을 릴리스하지 않는다.",
        "이 문서는 Gerber, BOM, CPL, 생산 주문 패키지를 릴리스하지 않는다.",
        "이 문서는 펌웨어 포트를 구현하지 않는다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "참고 자료", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_source_link(p, "기존 프로젝트", "https://github.com/Z3r0c0k3/hacking-box")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_source_link(p, "아트워크 소스", "black-white-ring.svg, provided Aegis ring artwork source")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_source_link(p, "로컬 계획 문서", "docs/rev3/design/hardware-spec.md, docs/rev3/design/decisions.md, docs/rev3/design/pin-map.md")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
