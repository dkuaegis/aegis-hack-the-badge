from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "rev3" / "sponsorship" / "jlcpcb" / "hack-the-badge-jlcpcb-support-brief-ko.docx"
PAGE_DIR = ROOT / "tmp" / "docx_assets" / "ko_pages"
LOGO = ROOT / "assets" / "brand" / "aegis" / "black-white-ring.png"
PCB_PREVIEW = ROOT / "hardware" / "releases" / "rev3" / "jlcpcb" / "preview" / "final_preview_iso.png"

FONT = "/System/Library/Fonts/Supplemental/AppleGothic.ttf"
W, H = 1650, 2134
M = 150
BLUE = (46, 116, 181)
DARK = (35, 35, 35)
MUTED = (92, 92, 92)
GRID = (183, 196, 212)
FILL = (232, 238, 245)
FILL2 = (242, 244, 247)


def font(size):
    return ImageFont.truetype(FONT, size)


F_TITLE = font(58)
F_H1 = font(38)
F_H2 = font(30)
F_BODY = font(25)
F_SMALL = font(21)
F_TINY = font(18)
F_BOLD = font(25)


def text_w(draw, text, f):
    return draw.textbbox((0, 0), text, font=f)[2]


def wrap_text(draw, text, f, max_width):
    lines = []
    for para in text.split("\n"):
        words = para.split(" ")
        current = ""
        for word in words:
            test = word if not current else current + " " + word
            if text_w(draw, test, f) <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                if text_w(draw, word, f) <= max_width:
                    current = word
                else:
                    chunks = textwrap.wrap(word, width=18)
                    lines.extend(chunks[:-1])
                    current = chunks[-1] if chunks else ""
        if current:
            lines.append(current)
    return lines


def draw_para(draw, xy, text, f=F_BODY, fill=DARK, max_width=None, line_gap=10):
    x, y = xy
    max_width = max_width or (W - 2 * M)
    for line in wrap_text(draw, text, f, max_width):
        draw.text((x, y), line, font=f, fill=fill)
        y += f.size + line_gap
    return y


def draw_bullets(draw, x, y, items, f=F_BODY, max_width=None):
    max_width = max_width or (W - 2 * M - 40)
    for item in items:
        draw.text((x, y + 3), "•", font=f, fill=DARK)
        y = draw_para(draw, (x + 45, y), item, f=f, max_width=max_width)
        y += 6
    return y


def new_page():
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    draw.text((M, 82), "Hack The Badge / Hacking Box V2 - Korean Review Copy", font=F_TINY, fill=MUTED)
    draw.text((W - M - 260, H - 82), "Phase 1 requirements draft", font=F_TINY, fill=MUTED)
    return img, draw


def draw_heading(draw, y, text):
    draw.text((M, y), text, font=F_H1, fill=BLUE)
    return y + 58


def draw_subheading(draw, y, text):
    draw.text((M, y), text, font=F_H2, fill=BLUE)
    return y + 44


def draw_table(draw, x, y, headers, rows, widths, row_h=72):
    header_h = 62
    cx = x
    for i, h in enumerate(headers):
        draw.rectangle((cx, y, cx + widths[i], y + header_h), fill=FILL, outline=GRID, width=2)
        draw.text((cx + 16, y + 16), h, font=F_SMALL, fill=DARK)
        cx += widths[i]
    y += header_h
    for row in rows:
        heights = []
        for i, cell in enumerate(row):
            lines = wrap_text(draw, cell, F_TINY, widths[i] - 28)
            heights.append(max(row_h, len(lines) * 24 + 26))
        rh = max(heights)
        cx = x
        for i, cell in enumerate(row):
            draw.rectangle((cx, y, cx + widths[i], y + rh), fill="white", outline=GRID, width=2)
            ty = y + 14
            for line in wrap_text(draw, cell, F_TINY, widths[i] - 28):
                draw.text((cx + 14, ty), line, font=F_TINY, fill=DARK)
                ty += 24
            cx += widths[i]
        y += rh
    return y


def paste_center(page, path, y, width):
    img = Image.open(path).convert("RGBA")
    ratio = width / img.width
    size = (int(width), int(img.height * ratio))
    img = img.resize(size, Image.LANCZOS)
    x = (W - size[0]) // 2
    page.paste(Image.new("RGB", size, "white"), (x, y))
    page.paste(img, (x, y), img if img.mode == "RGBA" else None)
    return y + size[1]


def page1():
    img, d = new_page()
    y = 180
    d.text((M, y), "JLCPCB SUPPORT BRIEF - KOREAN REVIEW COPY", font=F_SMALL, fill=BLUE)
    y += 58
    d.text((M, y), "Hack The Badge / Hacking Box V2", font=F_TITLE, fill=DARK)
    y += 78
    y = draw_para(d, (M, y), "기존 Arduino Hacking Badge 프로젝트를 기반으로 한 Phase 1 하드웨어 요구사항", F_BODY, MUTED)
    y += 18
    meta = [
        ("용도", "JLCPCB 제작 및 PCBA 지원 검토용 문서의 한국어 검토본"),
        ("작성 주체", "Aegis / MSG CTF project team"),
        ("작성일", "2026년 8월 4일"),
        ("상태", "요구사항 초안; Phase 1에서는 회로도와 PCB 레이아웃을 만들지 않음"),
        ("기반 프로젝트", "https://github.com/Z3r0c0k3/hacking-box"),
    ]
    for k, v in meta:
        d.text((M, y), f"{k}: {v}", font=F_SMALL, fill=DARK)
        y += 36
    d.line((M, y + 32, W - M, y + 32), fill=BLUE, width=5)
    y += 82
    y = draw_heading(d, y, "요약")
    y = draw_para(d, (M, y), "Hack The Badge는 Aegis / MSG CTF 행사 부스에서 사용할 CTF 하드웨어 해킹 배지다. 이번 JLCPCB 지원 요청의 목적은 회로도 작성 전에 하드웨어 콘셉트를 검토하고, PCBA 가능성과 제조/부품 선정 리스크를 미리 확인하는 것이다.")
    y += 14
    y = draw_para(d, (M, y), "이 V2 작업은 기존 Z3r0c0k3/hacking-box Arduino Hacking Badge 프로젝트를 기반으로 한다. 기존 프로젝트는 Serial shell 형태의 챌린지, 여러 독립 퍼즐, LED 피드백, solve 상태 저장을 제공한다.")
    y += 50
    y = draw_subheading(d, y, "예시 이미지")
    y = paste_center(img, LOGO, y, 360)
    d.text((M + 300, y + 20), "그림 1. black-white-ring.svg에서 렌더링한 Aegis 링 아트워크.", font=F_TINY, fill=MUTED)
    return img


def page2():
    img, d = new_page()
    y = 170
    y = paste_center(img, PCB_PREVIEW, y, 720)
    d.text((M + 70, y + 18), "그림 2. 기존 로컬 배지 렌더 이미지. 시각 방향 참고용이며 새 PCB 릴리스가 아님.", font=F_TINY, fill=MUTED)
    y += 76
    y = draw_heading(d, y, "기존 Hacking Box 프로젝트 기반")
    y = draw_para(d, (M, y), "공개 기준 프로젝트는 GitHub의 Z3r0c0k3/hacking-box다. README는 Arduino Hacking Badge를 설명하며, 대형 CTF 결승에서 배포되는 badge hacking 경험을 모방하는 구성을 가진다.")
    y += 12
    y = draw_bullets(d, M, y, [
        "기존 프로젝트는 4개의 독립 퍼즐 슬롯을 설명한다.",
        "각 퍼즐 solve 상태는 LED 피드백으로 표시되며, 모든 퍼즐을 풀었을 때 추가 LED 애니메이션을 수행한다.",
        "Serial Monitor 기반 shell 인터페이스로 게임 시작, 문제 선택, 상태 확인, 진행 상황 초기화, 화면 지우기, 문제 종료를 수행한다.",
        "README는 4800 baud Serial Monitor 설정과 Arduino Uno/Nano 스타일 하드웨어 기준을 문서화한다.",
        "긴 문제 텍스트와 ASCII art는 PROGMEM에 저장하고, solve 상태는 EEPROM에 저장한다.",
    ])
    y += 25
    y = draw_heading(d, y, "JLCPCB에 요청하려는 지원")
    draw_bullets(d, M, y, [
        "회로도와 PCB 레이아웃 전에 제안 하드웨어 구조를 검토한다.",
        "소량 SMT 조립에 적합한 JLCPCB 수급 가능 부품을 추천받는다.",
        "OLED 모듈을 JLCPCB PCBA에 포함할지, 후납땜으로 처리할지 조언받는다.",
        "USB-C, ESD, 입력 보호, 노출 UART/GPIO 패드, 버튼, LED, 선택 부저가 PCBA에 적합한지 검토받는다.",
    ])
    return img


def page3():
    img, d = new_page()
    y = 170
    y = draw_heading(d, y, "목표 하드웨어 기준안")
    y = draw_table(d, M, y, ["기능", "Phase 1 기준안", "이유"], [
        ("MCU", "ESP32-S3-WROOM 모듈", "Native USB, 선택적 Wi-Fi/BLE, 더 큰 메모리, 낮은 RF 리스크."),
        ("로직 전압", "3.3 V", "ESP32-S3와 최신 주변 모듈 기준에 맞음."),
        ("USB", "USB-C device port", "전원, Serial console, 펌웨어 업로드, 운영진 복구에 사용."),
        ("디스플레이", "0.96 inch 128x64 I2C OLED", "로고, 상태, 문제 텍스트를 작게 표시 가능."),
        ("입력", "택트 스위치 3개", "Left/OK/Right 조작, 미니게임 입력, 관리자 unlock 조합."),
        ("출력", "상태 LED 5개", "문제 진행 상태와 solve 여부를 명확히 표시."),
        ("오디오", "승인 시 드라이버가 있는 패시브 부저", "MCU GPIO에 직접 부하를 걸지 않고 효과음 제공."),
        ("챌린지 I/O", "UART와 보호된 GPIO 패드", "참가자를 위한 의도된 하드웨어 해킹 표면."),
    ], [250, 430, 670])
    y += 48
    y = draw_heading(d, y, "호환성 요구사항")
    y = draw_para(d, (M, y), "V2는 Arduino Nano / ATmega328P 기준 구현과 바이너리 호환되는 대체품이 아니라, ESP32-S3 기준 펌웨어 이식 또는 재작성으로 보아야 한다.")
    draw_bullets(d, M, y + 10, [
        "USB Serial shell 형태의 챌린지 흐름을 유지한다.",
        "가능한 범위에서 문제 선택 모델과 status 명령 흐름을 유지한다.",
        "solve 상태 피드백을 여러 LED에 매핑하고 all-solved 애니메이션 여지를 둔다.",
        "EEPROM 기반 상태 저장은 ESP32-S3 비휘발성 저장소로 대체한다.",
        "3.3 V 보드에서 신호를 노출하기 전에 V1의 5 V Arduino 의존성을 검토한다.",
        "USB flashing, EN/reset, BOOT 접근을 통해 운영진 현장 복구가 가능해야 한다.",
    ])
    return img


def page4():
    img, d = new_page()
    y = 170
    y = draw_heading(d, y, "제안 ESP32-S3 핀맵")
    y = draw_table(d, M, y, ["기능", "신호", "동작", "메모"], [
        ("USB D-", "GPIO19", "Native USB", "ESP32-S3 USB 사용 시 고정에 가까움."),
        ("USB D+", "GPIO20", "Native USB", "ESP32-S3 USB 사용 시 고정에 가까움."),
        ("OLED SDA", "GPIO4", "3.3 V I2C", "최종 OLED pull-up 여부 확인."),
        ("OLED SCL", "GPIO5", "3.3 V I2C", "확인 전까지 display address 0x3C 가정."),
        ("Challenge 0", "GPIO6", "직렬 저항을 거친 3.3 V", "챌린지 설계 승인 후 노출."),
        ("Challenge 1", "GPIO7", "직렬 저항을 거친 3.3 V", "챌린지 설계 승인 후 노출."),
        ("Challenge 2", "GPIO8", "직렬 저항을 거친 3.3 V", "챌린지 설계 승인 후 노출."),
        ("Left / OK / Right", "GPIO9 / GPIO10 / GPIO12", "Active-low inputs", "Pull-up과 debounce 적용."),
        ("Status LEDs", "GPIO13-GPIO17", "Active-high outputs", "부팅 시 low로 초기화."),
        ("Buzzer PWM", "GPIO18", "MOSFET-driven PWM", "부저 승인 시에만 사용."),
        ("Player UART TX/RX", "GPIO43 / GPIO44", "직렬 저항을 거친 3.3 V", "5 V tolerant 아님."),
        ("Boot / reset", "GPIO0 / EN", "운영진 복구", "챌린지 GPIO로 사용하지 않음."),
    ], [250, 260, 390, 450], row_h=66)
    y += 45
    y = draw_heading(d, y, "회로도 작성 전 승인 필요 항목")
    draw_table(d, M, y, ["ID", "필요한 결정", "중요한 이유"], [
        ("O-001", "정확한 ESP32-S3-WROOM 모듈 variant 확정.", "BOM, 펌웨어 파티션, 수급성, 가격에 영향."),
        ("O-002", "첫 제작 수량 확정.", "단가, 예비품 전략, 조립 옵션에 영향."),
        ("O-003", "OLED를 PCBA에 포함할지 후납땜할지 결정.", "BOM/CPL, footprint, 교체 workflow에 영향."),
        ("O-004", "부저를 포함할지 제거할지 승인.", "피드백은 좋지만 보드 면적과 펌웨어 작업 증가."),
        ("O-005", "챌린지 패드 동작과 라벨 확정.", "핵심 하드웨어 해킹 경험을 정의."),
    ], [150, 570, 630], row_h=64)
    return img


def page5():
    img, d = new_page()
    y = 170
    y = draw_heading(d, y, "JLCPCB 검토 질문")
    y = draw_bullets(d, M, y, [
        "현재 PCBA 수급성과 비용 기준으로 어떤 ESP32-S3-WROOM 모듈 variant를 추천하는가?",
        "OLED 모듈을 PCBA에 포함하는 것이 좋은가, 아니면 수동 조립으로 처리하는 것이 좋은가?",
        "제안한 USB-C receptacle, tactile switch, LED, passive buzzer가 단면 SMT 조립에 적합한가?",
        "참가자가 직접 만지는 USB 전원 보드에 input fuse/PTC를 넣는 것이 권장되는가?",
        "노출 UART와 challenge pad에 추가 ESD 또는 보호 부품이 필요한가?",
        "방패형 외형과 black/white 장식 silkscreen artwork에 DFM 우려가 있는가?",
    ])
    y += 36
    y = draw_heading(d, y, "Phase 1 범위 경계")
    y = draw_para(d, (M, y), "이 문서는 의도적으로 회로도 작성 전 단계에서 멈춘다. 다음 단계는 미결정 사항이 승인되고, V1 펌웨어 동작이 제안 ESP32-S3 핀맵과 대조된 뒤 시작하는 것이 적절하다.")
    y = draw_bullets(d, M, y + 8, [
        "이 문서는 새 회로도를 릴리스하지 않는다.",
        "이 문서는 새 PCB 레이아웃을 릴리스하지 않는다.",
        "이 문서는 Gerber, BOM, CPL, 생산 주문 패키지를 릴리스하지 않는다.",
        "이 문서는 펌웨어 포트를 구현하지 않는다.",
    ])
    y += 38
    y = draw_heading(d, y, "참고 자료")
    d.text((M, y), "기존 프로젝트: https://github.com/Z3r0c0k3/hacking-box", font=F_BODY, fill=DARK)
    y += 42
    d.text((M, y), "아트워크 소스: black-white-ring.svg, provided Aegis ring artwork source", font=F_BODY, fill=DARK)
    y += 42
    d.text((M, y), "로컬 계획 문서: docs/rev3/design/hardware-spec.md, docs/rev3/design/decisions.md, docs/rev3/design/pin-map.md", font=F_BODY, fill=DARK)
    return img


def build_docx(page_paths):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.2)
    sec.bottom_margin = Inches(0.2)
    sec.left_margin = Inches(0.25)
    sec.right_margin = Inches(0.25)
    for i, path in enumerate(page_paths):
        if i:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.add_run().add_picture(str(path), width=Inches(8.0))
    doc.save(OUT)


def main():
    PAGE_DIR.mkdir(parents=True, exist_ok=True)
    pages = [page1(), page2(), page3(), page4(), page5()]
    paths = []
    for idx, img in enumerate(pages, 1):
        path = PAGE_DIR / f"page-{idx}.png"
        img.save(path)
        paths.append(path)
    build_docx(paths)
    print(OUT)


if __name__ == "__main__":
    main()
