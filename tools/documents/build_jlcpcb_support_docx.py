from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SPONSORSHIP_DIR = ROOT / "docs" / "rev3" / "sponsorship" / "jlcpcb"
OUT = SPONSORSHIP_DIR / "hack-the-badge-jlcpcb-support-brief.docx"
LOGO = ROOT / "assets" / "brand" / "aegis" / "black-white-ring.png"
PCB_PREVIEW = (
    ROOT
    / "hardware"
    / "releases"
    / "rev3"
    / "jlcpcb"
    / "preview"
    / "final_preview_iso.png"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
MID_FILL = "E8EEF5"
BORDER = "B7C4D4"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="2E74B5", size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    font_name = "Arial Unicode MS" if name == "Calibri" else name
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_font(run, size=11, color=INK)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_fill(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for r in p.runs:
                    set_font(r, size=9.2, color=INK)
    set_table_width(table, widths)
    set_table_borders(table)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED, italic=True)
    return p


def add_source_link(paragraph, label, url):
    run = paragraph.add_run(label)
    set_font(run, size=10.5, color=INK, bold=True)
    paragraph.add_run(": ")
    link_run = paragraph.add_run(url)
    set_font(link_run, size=10.5, color=BLUE)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("JLCPCB SUPPORT BRIEF")
    set_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Hack The Badge / Hacking Box V2")
    set_font(r, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Phase 1 hardware requirements based on the existing Arduino Hacking Badge project")
    set_font(r, size=12, color=MUTED)

    meta = [
        ("Prepared for", "JLCPCB fabrication and PCBA support review"),
        ("Prepared by", "Aegis / MSG CTF project team"),
        ("Date", "August 4, 2026"),
        ("Status", "Requirements draft; schematic and PCB layout are out of scope for Phase 1"),
        ("Reference project", "https://github.com/Z3r0c0k3/hacking-box"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_font(vr, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule)


def add_image_block(doc):
    add_heading(doc, "Visual Reference", 2)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.9))
        add_caption(
            doc,
            "Figure 1. Example Aegis ring artwork rendered from black-white-ring.svg for badge identity reference.",
        )
    if PCB_PREVIEW.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(PCB_PREVIEW), width=Inches(3.7))
        add_caption(
            doc,
            "Figure 2. Existing local badge render used as a visual direction reference only; this brief does not release a new PCB design.",
        )


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "Hack The Badge / Hacking Box V2 - JLCPCB Support Brief"
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Phase 1 requirements draft"
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)

    add_masthead(doc)

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "Hack The Badge is a CTF booth hardware hacking badge for the Aegis / MSG CTF event. "
        "The requested JLCPCB support is focused on reviewing the hardware concept before schematic capture, "
        "confirming PCBA feasibility, and identifying manufacturability or component-selection risks early.",
    )
    add_para(
        doc,
        "This V2 effort is based on the existing Z3r0c0k3/hacking-box Arduino Hacking Badge project. "
        "The original project provides a serial shell-style challenge, multiple independent puzzles, LED feedback, "
        "and persistent solve state. V2 is intended to preserve the participant-visible challenge flow while moving "
        "to a repeatable ESP32-S3-based PCB suitable for a small event run.",
    )

    add_image_block(doc)

    add_heading(doc, "Based On The Existing Hacking Box Project", 1)
    add_para(
        doc,
        "The public reference project is Z3r0c0k3/hacking-box on GitHub. Its README describes an Arduino Hacking "
        "Badge that mimics the kind of badge hacking experience distributed at larger CTF finals. The repository "
        "contains example, old, and template badge firmware folders, and the README identifies Arduino IDE, C/C++, "
        "EEPROM, and PROGMEM as the technical foundation.",
    )
    for item in [
        "Four independent puzzle slots are described in the reference project.",
        "Each solved puzzle maps to LED feedback, with additional LED animation when all puzzles are solved.",
        "A shell-like serial monitor interface lets participants start the game, select problems, check status, reset progress, clear the screen, and exit a problem.",
        "The README documents 4800 baud serial monitor setup and an Arduino Uno/Nano-style hardware baseline.",
        "Large challenge text and ASCII art are stored in program memory using PROGMEM, while solve state is persisted using EEPROM.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Support Requested From JLCPCB", 1)
    for item in [
        "Review the proposed architecture before schematic capture and PCB layout.",
        "Recommend JLCPCB-available parts for low-volume SMT assembly.",
        "Advise whether the OLED module should be assembled by JLCPCB or hand-soldered after PCBA.",
        "Review USB-C, ESD, input protection, exposed UART/GPIO pads, buttons, LEDs, and optional buzzer for PCBA practicality.",
        "Identify DFM issues related to a badge-shaped PCB outline and decorative silkscreen artwork.",
        "Help confirm which decisions must be frozen before Gerber, BOM, and CPL generation.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Target Hardware Baseline", 1)
    add_table(
        doc,
        ["Function", "Phase 1 baseline", "Reason"],
        [
            ("MCU", "ESP32-S3-WROOM module", "Native USB, optional Wi-Fi/BLE, more memory, and lower RF risk than a bare chip."),
            ("Logic voltage", "3.3 V", "Matches ESP32-S3 and modern peripheral modules."),
            ("USB", "USB-C device port", "Power, serial console, firmware upload, and staff recovery."),
            ("Display", "0.96 inch 128x64 I2C OLED", "Compact logo, status, and challenge text display."),
            ("Input", "Three tactile switches", "Left/OK/right controls, mini-game input, or admin-unlock gesture."),
            ("Output", "Five status LEDs", "Clear challenge progress and solved-state indication."),
            ("Audio", "Passive buzzer with driver, if approved", "Optional audible feedback without directly loading MCU GPIO."),
            ("Challenge I/O", "UART plus protected GPIO pads", "Intentional hardware hacking surface for participants."),
        ],
        [1800, 2700, 4860],
        MID_FILL,
    )

    add_heading(doc, "Compatibility Requirements", 1)
    add_para(
        doc,
        "V2 should be treated as a firmware port or rewrite, not as a direct binary-compatible replacement for the "
        "Arduino Nano / ATmega328P reference implementation. The participant-visible behavior should remain familiar.",
    )
    for item in [
        "Preserve the USB Serial shell-style challenge workflow.",
        "Preserve the puzzle selection model and status commands where practical.",
        "Map solved-state feedback to multiple LEDs, with room for an all-solved animation.",
        "Replace EEPROM persistence with ESP32-S3 non-volatile storage while preserving visible behavior.",
        "Review any V1 dependency on 5 V Arduino electrical behavior before exposing signals on a 3.3 V board.",
        "Keep staff recovery practical through USB flashing plus EN/reset and BOOT access.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Proposed ESP32-S3 Pin Map", 1)
    add_table(
        doc,
        ["Function", "Signal", "Behavior", "Notes"],
        [
            ("USB D-", "GPIO19", "Native USB", "Fixed for ESP32-S3 USB use."),
            ("USB D+", "GPIO20", "Native USB", "Fixed for ESP32-S3 USB use."),
            ("OLED SDA", "GPIO4", "3.3 V I2C", "Confirm final OLED pull-ups."),
            ("OLED SCL", "GPIO5", "3.3 V I2C", "Assume display address 0x3C until confirmed."),
            ("Challenge 0", "GPIO6", "3.3 V through series resistor", "Expose only after challenge approval."),
            ("Challenge 1", "GPIO7", "3.3 V through series resistor", "Expose only after challenge approval."),
            ("Challenge 2", "GPIO8", "3.3 V through series resistor", "Expose only after challenge approval."),
            ("Left / OK / Right", "GPIO9 / GPIO10 / GPIO12", "Active-low inputs", "Use pull-ups and debounce."),
            ("Status LEDs", "GPIO13-GPIO17", "Active-high outputs", "Initialize low at boot."),
            ("Buzzer PWM", "GPIO18", "MOSFET-driven PWM", "Only if buzzer is approved."),
            ("Player UART TX/RX", "GPIO43 / GPIO44", "3.3 V through series resistors", "Not 5 V tolerant."),
            ("Boot / reset", "GPIO0 / EN", "Staff recovery", "Do not use as challenge GPIO."),
        ],
        [1800, 1800, 2760, 3000],
        LIGHT_FILL,
    )

    add_heading(doc, "Open Decisions Before Schematic Capture", 1)
    add_table(
        doc,
        ["ID", "Decision needed", "Why it matters"],
        [
            ("O-001", "Confirm exact ESP32-S3-WROOM module variant.", "Affects BOM, firmware partitioning, availability, and price."),
            ("O-002", "Confirm first-run quantity.", "Affects unit price, spare strategy, and assembly options."),
            ("O-003", "Decide whether OLED is PCBA-installed or hand-soldered.", "Affects BOM/CPL, footprint, and replacement workflow."),
            ("O-004", "Approve or remove the buzzer.", "Adds feedback but increases board area and firmware work."),
            ("O-005", "Finalize challenge pad behavior and labels.", "Defines the core hardware hacking experience."),
            ("O-006", "Define Wi-Fi/BLE management policy.", "Management access must not become an unintended bypass."),
            ("O-007", "Decide protection level for USB and exposed GPIO/UART.", "Participant handling increases electrical abuse risk."),
            ("O-008", "Approve badge shape and silkscreen artwork.", "Impacts DFM, branding, and support review."),
        ],
        [1000, 3900, 4460],
        MID_FILL,
    )

    add_heading(doc, "Questions For JLCPCB Review", 1)
    for item in [
        "Which ESP32-S3-WROOM module variant is recommended for current PCBA availability and cost?",
        "Should the OLED module be included in PCBA or handled as manual assembly?",
        "Are the proposed USB-C receptacle, tactile switches, LEDs, and passive buzzer suitable for single-side SMT assembly?",
        "Is an input fuse/PTC recommended for a participant-handled USB-powered board?",
        "Are additional ESD or protection components recommended for exposed UART and challenge pads?",
        "Are there DFM concerns with a shield-shaped outline and decorative black/white silkscreen artwork?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Phase 1 Scope Boundary", 1)
    add_para(
        doc,
        "This brief intentionally stops before schematic capture. The next phase should begin only after the open decisions "
        "above are approved and the V1 firmware behavior has been reviewed against the proposed ESP32-S3 pin map.",
    )
    for item in [
        "No new schematic is released by this document.",
        "No new PCB layout is released by this document.",
        "No Gerber, BOM, CPL, or production order package is released by this document.",
        "No firmware port is implemented by this document.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "References", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_source_link(p, "Existing project", "https://github.com/Z3r0c0k3/hacking-box")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_source_link(p, "Artwork source", "black-white-ring.svg, provided Aegis ring artwork source")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_source_link(p, "Local planning documents", "docs/rev3/design/hardware-spec.md, docs/rev3/design/decisions.md, docs/rev3/design/pin-map.md")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
